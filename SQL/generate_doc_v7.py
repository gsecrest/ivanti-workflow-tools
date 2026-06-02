from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading(text, level=1):
    return doc.add_heading(text, level=level)

def body(text):
    return doc.add_paragraph(text)

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0x7C)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'EEF2F8')
    p._p.get_or_add_pPr().append(shading)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1F3864')
        hdr[i]._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_heading('FindTeamByBlockTypeAndWorkflow_v7.sql', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Technical Documentation — Ivanti ITSM Workflow Analysis Query (Optimised)')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
sub.runs[0].font.italic = True
doc.add_paragraph()

# ── 1. Overview ───────────────────────────────────────────────────────────────
heading('1. Overview')
body(
    'This SQL script is a diagnostic and audit tool for the Ivanti ITSM (Neurons) platform. '
    'It inspects published workflow definitions and identifies which teams are assigned within '
    'specific workflow blocks — either via Quick Actions (update/advancedtask blocks) or direct '
    'task block assignments. The results help administrators understand team routing logic '
    'embedded in service request workflows without navigating the Ivanti UI.'
)
body(
    'v7 builds on v6 with three targeted improvements: WITH (NOLOCK) hints on all base-table '
    'reads to avoid blocking the live Ivanti system, removal of DISTINCT from the #AllBlocks '
    'XML shred stage (the XML column type is not comparable in SQL Server; deduplication '
    'is handled correctly in the downstream PATH steps), and a two-step null guard in the '
    'CHARINDEX-based OwnerTeam extraction that prevents returning arbitrary text when a '
    'Quick Action\'s ExpressionText is null. All output columns and result sets are identical '
    'to v6.'
)

# ── 2. Input Parameters ───────────────────────────────────────────────────────
heading('2. Input Parameters')
body('Four DECLARE variables at the top of the script control filtering. Leave any blank to return all matching records.')
doc.add_paragraph()
add_table(
    ['Parameter', 'Type', 'Default Value', 'Description'],
    [
        ['@WorkflowName', 'NVARCHAR(255)', '(blank — all)', 'Filter by workflow name (partial match).'],
        ['@BlockType',    'NVARCHAR(50)',  '(blank — all)', 'Filter by block type: task, advancedtask, or update.'],
        ['@TeamName',     'NVARCHAR(255)', '(blank — all)', 'Filter by team name (partial match).'],
        ['@Status',       'NVARCHAR(50)',  '(blank — all)', 'Filter by request offering status (e.g. Published, Design).'],
    ]
)

# ── 3. How It Works ───────────────────────────────────────────────────────────
heading('3. How It Works')
body('The script executes in five logical stages followed by a final UNION result:')

heading('Stage 1 — Load Filtered Workflows  (#FilteredWorkflows)', level=2)
body(
    'A CTE (LatestVersions) uses ROW_NUMBER() partitioned by WorkflowTypeLink_RecID to identify '
    'the most recent version of each workflow in a single pass. WITH (NOLOCK) is applied to '
    'frs_def_workflow_definition and frs_def_workflow_type — this is a read-only reporting query '
    'and NOLOCK prevents it from blocking or being blocked by concurrent writes on the live '
    'Ivanti system. The workflow XML is cleaned of its declaration and namespace prefix, cast '
    'to the XML data type, and stored alongside a UPPER()-normalised RecID. A clustered index '
    'on WorkflowDefinitionRecID is created immediately after population.'
)
code_block(
    ";WITH LatestVersions AS (\n"
    "    SELECT RecID, WorkflowTypeLink_RecID, DefVersion, Details,\n"
    "           ROW_NUMBER() OVER (\n"
    "               PARTITION BY WorkflowTypeLink_RecID\n"
    "               ORDER BY CAST(DefVersion AS INT) DESC\n"
    "           ) AS rn\n"
    "    FROM frs_def_workflow_definition WITH (NOLOCK)\n"
    ")\n"
    "SELECT ... INTO #FilteredWorkflows\n"
    "FROM LatestVersions lv\n"
    "JOIN frs_def_workflow_type wt WITH (NOLOCK) ON ...\n"
    "WHERE lv.rn = 1 AND wt.Name LIKE '%form' ...;\n"
    "CREATE CLUSTERED INDEX IX_FW_RecID ON #FilteredWorkflows (WorkflowDefinitionRecID);"
)

heading('Stage 2 — Single XML Shred Pass  (#AllBlocks)', level=2)
body(
    'All workflow XML is shredded into individual blocks in a single pass using '
    'CROSS APPLY ... .nodes(). The block fragment (block.query(\'.\')) is stored alongside '
    'the title and type so PATH-specific stages below can read the cheap stored fragment '
    'rather than re-shredding the full document. DISTINCT is intentionally omitted here: '
    'the SQL XML type is not comparable in SQL Server, so DISTINCT on a column containing '
    'XML would produce an error. Deduplication is handled in the PATH steps which do not '
    'carry the XML column forward.'
)
code_block(
    "-- No DISTINCT: XML type is not comparable; dedup happens at the PATH steps.\n"
    "SELECT\n"
    "    fw.WorkflowName, fw.WorkflowDefinitionRecID, fw.DefVersion,\n"
    "    LTRIM(RTRIM(b.block.value('(title)[1]', 'nvarchar(255)'))) AS BlockTitle,\n"
    "    LTRIM(RTRIM(b.block.value('(type)[1]',  'nvarchar(50)')))  AS BlockType,\n"
    "    b.block.query('.')                                          AS BlockXml\n"
    "INTO #AllBlocks\n"
    "FROM #FilteredWorkflows fw\n"
    "CROSS APPLY fw.XmlData.nodes('/scenario/blocks/block') b(block);\n"
    "CREATE CLUSTERED INDEX IX_AB_RecID ON #AllBlocks (WorkflowDefinitionRecID);"
)

heading('Stage 3 — Extract Quick Action Blocks  (#Blocks)  [PATH 1]', level=2)
body(
    'Reads the stored block XML fragment from #AllBlocks and extracts blocks containing a '
    'QuickAction property. The QAID (Quick Action ID) is cast to uniqueidentifier so the '
    'later JOIN to frs_def_quick_actions.Id is type-matched and sargable. A clustered index '
    'on QAID is created after population.'
)
code_block(
    "SELECT DISTINCT ab.WorkflowName, ab.WorkflowDefinitionRecID, ab.DefVersion,\n"
    "    ab.BlockTitle, ab.BlockType,\n"
    "    TRY_CAST(\n"
    "        q.qaprop.value('(groups/group/param[name=\"QAID\"]/value)[1]', 'nvarchar(100)')\n"
    "    AS uniqueidentifier) AS QAID\n"
    "INTO #Blocks\n"
    "FROM #AllBlocks ab\n"
    "CROSS APPLY ab.BlockXml.nodes('block/blockProperties/property[name=\"QuickAction\"]') q(qaprop);\n"
    "CREATE CLUSTERED INDEX IX_Blocks_QAID ON #Blocks (QAID);"
)

heading('Stage 4 — Extract Task Blocks  (#TaskBlocks)  [PATH 2]', level=2)
body(
    'Reads the stored block XML from #AllBlocks and extracts blocks with a teamblock property. '
    'CROSS APPLY (VALUES) computes the TeamName expression once, avoiding repeated XPath '
    'evaluation. Blank values and dynamic expressions (starting with "$(") are filtered out. '
    'A clustered index on WorkflowDefinitionRecID is created after population.'
)
code_block(
    "SELECT DISTINCT ab.WorkflowName, ab.WorkflowDefinitionRecID, ab.DefVersion,\n"
    "    ab.BlockTitle, ab.BlockType, tv.TeamName\n"
    "INTO #TaskBlocks\n"
    "FROM #AllBlocks ab\n"
    "CROSS APPLY ab.BlockXml.nodes('block/blockProperties/property[name=\"teamblock\"]"
    "/groups/group/param[name=\"team\"]') p(prop)\n"
    "CROSS APPLY (VALUES (LTRIM(RTRIM(p.prop.value('(value)[1]', 'nvarchar(255)'))))) tv(TeamName)\n"
    "WHERE tv.TeamName <> '' AND tv.TeamName NOT LIKE '$(%' ...;\n"
    "CREATE CLUSTERED INDEX IX_TaskBlocks_RecID ON #TaskBlocks (WorkflowDefinitionRecID);"
)

heading('Stage 5 — Materialise Offering Status  (#WorkflowOffering)', level=2)
body(
    'The three-table offering status join (ServiceReqFulfillmentPlan → FusionLink → '
    'ServiceReqTemplate) is materialised once into a temp table with a clustered index. '
    'Both UNION branches in the final SELECT read from this table rather than re-executing '
    'the join. WITH (NOLOCK) is applied to all three base tables.'
)
code_block(
    "SELECT DISTINCT UPPER(fp.WorkflowId) AS WorkflowId, srt.Status\n"
    "INTO #WorkflowOffering\n"
    "FROM ServiceReqFulfillmentPlan fp WITH (NOLOCK)\n"
    "JOIN FusionLink fl WITH (NOLOCK)\n"
    "    ON fl.TargetID = fp.RecId AND fl.RelationshipName = 'ServiceReqTemplate...'\n"
    "JOIN ServiceReqTemplate srt WITH (NOLOCK) ON srt.RecId = fl.SourceID;\n"
    "CREATE CLUSTERED INDEX IX_WO_WorkflowId ON #WorkflowOffering (WorkflowId);"
)

heading('Final UNION Result', level=2)
body(
    'UNIONs two branches, both joined to #WorkflowOffering for the status column:'
)
body(
    '#Blocks branch (Quick Action blocks): joins to frs_def_quick_actions WITH (NOLOCK) and '
    'uses a two-step CHARINDEX scan to extract the team name from the Definition JSON column. '
    'OPENJSON cannot be used here because Ivanti stores JavaScript Date literals '
    '(new Date(...)) in this column, which are not valid JSON — OPENJSON validates the full '
    'document and rejects it before reaching $.FieldValues. CHARINDEX tolerates the '
    'non-standard format.'
)
body(
    'The improved two-step null guard in v7: ownerPos finds "FieldName":"OwnerTeam"; '
    'etPos searches forward for "ExpressionText":" only when ownerPos.pos > 0. If '
    'ExpressionText is null in the JSON ("ExpressionText":null), the second CHARINDEX '
    'returns 0, etPos.pos = 0, and TeamName is correctly set to NULL rather than '
    'grabbing arbitrary text from the string.'
)
code_block(
    "CROSS APPLY (VALUES (\n"
    "    CHARINDEX('\"FieldName\":\"OwnerTeam\"', qa.Definition)\n"
    ")) ownerPos (pos)\n"
    "CROSS APPLY (VALUES (\n"
    "    CASE WHEN ownerPos.pos > 0\n"
    "         THEN CHARINDEX('\"ExpressionText\":\"', qa.Definition, ownerPos.pos)\n"
    "         ELSE 0\n"
    "    END\n"
    ")) etPos (pos)\n"
    "CROSS APPLY (VALUES (\n"
    "    CASE WHEN etPos.pos > 0\n"
    "         THEN LEFT(SUBSTRING(qa.Definition, etPos.pos + 18, 500),\n"
    "                   CHARINDEX('\"', SUBSTRING(qa.Definition, etPos.pos + 18, 500)) - 1)\n"
    "    END\n"
    ")) tn (TeamName)"
)
body(
    '#TaskBlocks branch (task blocks): team name was already extracted from XML in Stage 4; '
    'no further lookup is needed. Results from both branches are sorted by WorkflowName, '
    'BlockType, BlockTitle.'
)

# ── 4. Changes from v6 to v7 ─────────────────────────────────────────────────
heading('4. Changes from v6 to v7')
add_table(
    ['#', 'Change', 'Reason', 'Benefit'],
    [
        ['1',
         'WITH (NOLOCK) added to all base-table reads (frs_def_workflow_definition, '
         'frs_def_workflow_type, ServiceReqFulfillmentPlan, FusionLink, ServiceReqTemplate, '
         'frs_def_quick_actions)',
         'Read-only reporting query running against a live Ivanti system',
         'Prevents the query from blocking or being blocked by concurrent writes'],
        ['2',
         'DISTINCT removed from #AllBlocks (Stage 2)',
         'SQL Server XML type is not comparable — DISTINCT on an XML column would error',
         'Eliminates a latent bug; deduplication still occurs correctly in the PATH steps'],
        ['3',
         'Two-step null guard in CHARINDEX OwnerTeam extraction: checks etPos.pos > 0 '
         'before extracting ExpressionText',
         'When ExpressionText is null, the original single-step guard could return '
         'arbitrary text from the string instead of NULL',
         'Correctly returns NULL for Quick Action blocks where no team is assigned'],
        ['4',
         'All parameter defaults changed to blank (empty string)',
         'Blank defaults match all values — more appropriate for a general-purpose script',
         'Script returns all results out of the box without editing parameters first'],
    ]
)

# ── 5. Database Tables Referenced ────────────────────────────────────────────
heading('5. Database Tables Referenced')
add_table(
    ['Table', 'Purpose'],
    [
        ['frs_def_workflow_definition',  'Stores versioned workflow XML definitions'],
        ['frs_def_workflow_type',        'Stores workflow names and types'],
        ['frs_def_quick_actions',        'Stores Quick Action JSON definitions (contains OwnerTeam assignments)'],
        ['ServiceReqFulfillmentPlan',    'Links workflows to service request fulfillment plans'],
        ['FusionLink',                   'Relationship table linking fulfillment plans to service request templates'],
        ['ServiceReqTemplate',           'Service request offering metadata including Status'],
        ['StandardUserTeam',             'Active service desk teams (used by the web app dropdown)'],
    ]
)

# ── 6. Output Columns ─────────────────────────────────────────────────────────
heading('6. Output Columns')
add_table(
    ['Column', 'Description'],
    [
        ['WorkflowName',          'Name of the workflow (e.g. "Audit Data Request Request form")'],
        ['DefVersion',            'Version number of the workflow definition'],
        ['RequestOfferingStatus', 'Status of the associated request offering, or "No Offering" if unlinked'],
        ['BlockTitle',            'Display title of the block within the workflow'],
        ['BlockType',             'Type of block: task, advancedtask, or update'],
        ['TeamName',              'Team assigned within that block'],
    ]
)

# ── 7. Sample Results ─────────────────────────────────────────────────────────
heading('7. Sample Results  (Team = "Risk Management Support", Status = "Published")')
add_table(
    ['WorkflowName', 'Ver', 'Offering Status', 'Block Title', 'Block Type', 'Team'],
    [
        ['Audit Data Request Request form',                         '11', 'Published (Automatic)', 'Set Team and Classification',                'update', 'Risk Management Support'],
        ['Audit Data Request Request form',                         '11', 'Published (Automatic)', 'Set Team and Classification 2',              'update', 'Risk Management Support'],
        ['Request Software Not in the Company Portal Request form', '12', 'Published (Automatic)', '1 - Fulfill Risk Assessment Request Lite',  'task',   'Risk Management Support'],
        ['Request Software Not in the Company Portal Request form', '12', 'Published (Automatic)', '1a - Fulfill Risk Assessment Request',      'task',   'Risk Management Support'],
    ]
)

# ── 8. Common Use Cases ───────────────────────────────────────────────────────
heading('8. Common Use Cases')
bullet('Audit which workflows route work to a specific team before renaming or retiring that team.')
bullet('Identify all task and update blocks assigned to a team across all published workflows.')
bullet('Verify workflow version alignment — confirm the latest version is what is actively used.')
bullet('Pre-migration checklist — find all workflows referencing a team being restructured.')
bullet('Set @TeamName and @Status to blank for a full cross-team routing map.')

# ── 9. Tips & Customization ───────────────────────────────────────────────────
heading('9. Tips & Customization')
code_block(
    "DECLARE @WorkflowName NVARCHAR(255) = '';          -- blank = all workflows\n"
    "DECLARE @BlockType    NVARCHAR(50)  = 'task';      -- only task blocks\n"
    "DECLARE @TeamName     NVARCHAR(255) = '';          -- blank = all teams\n"
    "DECLARE @Status       NVARCHAR(50)  = 'Published'; -- only published offerings"
)
bullet("Change @BlockType to 'update' to see only Quick Action assignment blocks.")
bullet('Leave @TeamName blank to generate a full routing map across all teams.')
bullet("Change @Status to 'Design' to audit workflows still in development.")
bullet("Workflow filter targets names ending in 'form' — adjust LIKE '%form' to broaden scope.")

# ── 10. Dependencies & Notes ──────────────────────────────────────────────────
heading('10. Dependencies & Notes')
bullet('Requires SELECT access on: frs_def_workflow_definition, frs_def_workflow_type, frs_def_quick_actions, ServiceReqFulfillmentPlan, FusionLink, ServiceReqTemplate.')
bullet('Uses temporary tables and CREATE INDEX — requires tempdb write access.')
bullet('Temp tables are cleaned up at the end via DROP TABLE statements.')
bullet('XML namespace and declaration are stripped before parsing — required because SQL Server rejects the Ivanti-specific namespace URI.')
bullet(
    'Quick Action team extraction uses CHARINDEX string parsing on the Definition column. '
    'OPENJSON cannot be used because Ivanti stores JavaScript Date literals (new Date(...)) '
    'in that column, which are not valid JSON. CHARINDEX tolerates the non-standard format.'
)
bullet('Targets SQL Server 2016+ due to XQuery, TRY_CAST, and window function usage. Tested on SQL Server 2019.')

# ── 11. Web Application ───────────────────────────────────────────────────────
heading('11. Web Application')
body(
    'The Next.js web app (workflow-query-app) provides a browser-based interface over the '
    'v7 SQL query. It connects to the same SQL Server database and exposes the same four '
    'filter parameters through a form UI. The app is a full-stack TypeScript application '
    'using Next.js App Router with server-side API routes. Source code is at '
    'https://github.com/gsecrest/workflow-query-app.'
)

heading('11.1 Architecture', level=2)
body(
    'The app uses a shared connection pool (lib/db.ts) so a single persistent TCP connection '
    'is maintained across all requests rather than opening and closing a connection on every '
    'query. The SQL query is extracted to lib/workflow-query.ts and shared between the search '
    'and export routes, eliminating duplication.'
)
bullet('lib/db.ts — Shared mssql connection pool and DB config.')
bullet('lib/db-password.ts — Windows DPAPI password decryption utility (see Section 11.5).')
bullet('lib/workflow-query.ts — The v7 SQL query shared by /api/query and /api/export.')
bullet('app/api/query/route.ts — POST endpoint: accepts filters, returns JSON rows.')
bullet('app/api/teams/route.ts — GET endpoint: returns active service desk teams for the dropdown.')
bullet('app/api/export/workflow-results.csv/route.ts — GET endpoint: returns results as a CSV download.')
bullet('app/page.tsx — Single-page React UI with filter form and results table.')

heading('11.2 Filters', level=2)
body('The filter form maps directly to the four SQL parameters. All fields are optional — leaving any blank searches across all values.')
bullet('Workflow Name: Partial-match text input. Equivalent to @WorkflowName.')
bullet('Block Type: Dropdown: All block types, task, advancedtask, update. Equivalent to @BlockType.')
bullet('Team Name: Dropdown populated from active service desk teams in StandardUserTeam. Equivalent to @TeamName.')
bullet('Status: Dropdown: All statuses, Published, Design. Equivalent to @Status.')

heading('11.3 Results Table', level=2)
body(
    'Results render in a table with the same six columns returned by the SQL query. '
    'The results panel only appears after the first query runs.'
)

heading('11.4 Exporting Results', level=2)
body('Two export options appear in the results header whenever rows are returned:')
bullet(
    'Export CSV — Downloads the current result set as workflow-results.csv via a server-side '
    'route that re-runs the query. Values are double-quoted and a UTF-8 BOM is prepended so '
    'Excel opens the file directly without an import wizard. The filename is embedded in the '
    'URL path so managed Chrome browsers use it as the download name even when '
    'Content-Disposition headers are ignored.'
)
bullet(
    'Copy to Clipboard — Copies results as tab-separated values (TSV). Pasting directly into '
    'an open Excel sheet aligns data into columns without an import wizard.'
)

heading('11.5 Password Encryption (Windows DPAPI)', level=2)
body(
    'For production deployments on Windows, the database password can be encrypted using '
    'Windows DPAPI so it is not stored in plaintext. Set DB_PASSWORD_ENCRYPTED in .env.local '
    'instead of DB_PASSWORD. The encrypted value is tied to the Windows user account and '
    'machine that performed the encryption — it cannot be decrypted elsewhere.'
)
code_block(
    "# Encrypt in PowerShell:\n"
    "Add-Type -AssemblyName System.Security\n"
    "[System.Convert]::ToBase64String(\n"
    "  [System.Security.Cryptography.ProtectedData]::Protect(\n"
    "    [System.Text.Encoding]::UTF8.GetBytes('your-password'),\n"
    "    $null,\n"
    "    [System.Security.Cryptography.DataProtectionScope]::CurrentUser\n"
    "  )\n"
    ")"
)

heading('11.6 API Routes', level=2)
add_table(
    ['Route', 'Method', 'Description'],
    [
        ['/api/query',                         'POST', 'Accepts { workflowName, blockType, teamName, status } and returns JSON rows.'],
        ['/api/teams',                         'GET',  'Returns active service desk teams from StandardUserTeam for the dropdown.'],
        ['/api/export/workflow-results.csv',   'GET',  'Re-runs the query with filter params and returns a CSV download.'],
    ]
)

heading('11.7 Setup', level=2)
body('Requires Node.js 18+ and a .env.local file in the project root:')
code_block(
    "DB_SERVER=your-sql-server\n"
    "DB_DATABASE=your-database\n"
    "DB_USER=your-username\n"
    "DB_PASSWORD=your-password   # or DB_PASSWORD_ENCRYPTED for DPAPI\n"
    "DB_PORT=1433"
)
body('Start the development server:')
code_block(
    "npm install\n"
    "npm run dev"
)
body('The app runs at http://localhost:3000. For background deployment on Windows, run setup-windows.bat as Administrator to configure PM2.')

# ── 12. Repository Documentation ─────────────────────────────────────────────
heading('12. Repository Documentation')
body(
    'The ivanti-workflow-tools GitHub repository (https://github.com/gsecrest/ivanti-workflow-tools) '
    'contains the SQL scripts and supporting files. The workflow-query-app is maintained in a '
    'separate repository at https://github.com/gsecrest/workflow-query-app.'
)

heading('12.1 SQL Folder', level=2)
body('SQL/README.md is a detailed reference for every script in the SQL folder.')
body('Version history:')
add_table(
    ['File', 'Notes'],
    [
        ['FindTeamByBlockTypeAndWorkflow.sql',    'v1 — initial version'],
        ['FindTeamByBlockTypeAndWorkflow_v2.sql', 'v2'],
        ['FindTeamByBlockTypeAndWorkflow_v3.sql', 'v3'],
        ['FindTeamByBlockTypeAndWorkflow_v4.sql', 'v4 — see v4 documentation'],
        ['FindTeamByBlockTypeAndWorkflow_v5.sql', 'v5 — performance-optimised rewrite; same output as v4'],
        ['FindTeamByBlockTypeAndWorkflow_v6.sql', 'v6 — single XML shred pass (#AllBlocks); WorkflowOffering materialised as temp table'],
        ['FindTeamByBlockTypeAndWorkflow_v7.sql', 'v7 — NOLOCK hints, DISTINCT fix on #AllBlocks, improved CHARINDEX null guard'],
    ]
)

heading('12.2 Web App Documentation', level=2)
body('The workflow-query-app repository contains two documentation files:')
bullet('README.md — User-facing documentation covering what the app does, setup instructions, API route reference, project structure, troubleshooting, and a reference to the standalone SQL script.')
bullet('BUILDING.md — Step-by-step guide for building the app from scratch, covering each file, implementation decisions, and the SQL query walkthrough including why CHARINDEX is used instead of OPENJSON.')

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/Glenn.Secrest/Downloads/Claude/SQL/FindTeamByBlockTypeAndWorkflow_v7_Documentation.docx'
doc.save(out)
print(f'Saved: {out}')
