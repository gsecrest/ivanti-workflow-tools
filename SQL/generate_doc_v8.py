from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading(text, level=1):
    return doc.add_heading(text, level=level)

def body(text):
    return doc.add_paragraph(text)

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0x7C)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'EEF2F8')
    p._p.get_or_add_pPr().append(shading)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1F3864')
        hdr[i]._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_heading('FindTeamByBlockTypeAndWorkflow_v8.sql', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Technical Documentation — Ivanti ITSM Workflow Analysis Query (Optimised)')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
sub.runs[0].font.italic = True
doc.add_paragraph()

# ── 1. Overview ───────────────────────────────────────────────────────────────
heading('1. Overview')
body(
    'This SQL script is a diagnostic and audit tool for the Ivanti ITSM (Neurons) platform. '
    'It inspects published workflow definitions and identifies which teams or approval groups '
    'are assigned within specific workflow blocks. Results help administrators understand team '
    'routing and approval logic embedded in service request workflows without navigating the '
    'Ivanti UI.'
)
body(
    'v8 extends v7 with two additions: (1) PATH 3 for approval blocks (vote0007, vote) — '
    'extracts the approver contact group GUID from the workflow XML and joins to the '
    'ContactGroup table to resolve the name; (2) the QuickAction path (PATH 1) now explicitly '
    'covers create, notification, quickaction, and createnew0002 block types, which use the '
    'same QAID mechanism as advancedtask and update. All output columns are identical to v7.'
)

# ── 2. Input Parameters ───────────────────────────────────────────────────────
heading('2. Input Parameters')
body('Four DECLARE variables at the top of the script control filtering. Leave any blank to return all matching records.')
doc.add_paragraph()
add_table(
    ['Parameter', 'Type', 'Default Value', 'Description'],
    [
        ['@WorkflowName', 'NVARCHAR(255)', '(blank — all)', 'Filter by workflow name (partial match).'],
        ['@BlockType',    'NVARCHAR(50)',  '(blank — all)', 'Filter by block type (see supported types below).'],
        ['@TeamName',     'NVARCHAR(255)', '(blank — all)', 'Filter by team name or approval group name (partial match).'],
        ['@Status',       'NVARCHAR(50)',  '(blank — all)', 'Filter by request offering status (e.g. Published, Design).'],
    ]
)

# ── 3. Supported Block Types ──────────────────────────────────────────────────
heading('3. Supported Block Types')
add_table(
    ['Block Type', 'Path', 'Team Source'],
    [
        ['task',            'PATH 2', 'teamblock property in workflow XML'],
        ['advancedtask',    'PATH 1', 'OwnerTeam in frs_def_quick_actions.Definition via CHARINDEX'],
        ['update',          'PATH 1', 'OwnerTeam in frs_def_quick_actions.Definition via CHARINDEX'],
        ['create',          'PATH 1', 'OwnerTeam in frs_def_quick_actions.Definition via CHARINDEX'],
        ['notification',    'PATH 1', 'OwnerTeam in frs_def_quick_actions.Definition via CHARINDEX'],
        ['quickaction',     'PATH 1', 'OwnerTeam in frs_def_quick_actions.Definition via CHARINDEX'],
        ['createnew0002',   'PATH 1', 'OwnerTeam in frs_def_quick_actions.Definition via CHARINDEX'],
        ['vote0007',        'PATH 3', 'contactgroup GUID in workflow XML → ContactGroup.Name'],
        ['vote',            'PATH 3', 'contactgroup GUID in workflow XML → ContactGroup.Name'],
    ]
)

# ── 4. How It Works ───────────────────────────────────────────────────────────
heading('4. How It Works')
body('The script executes in three stages followed by a final three-way UNION result:')

heading('Stage 1 — Load Filtered Workflows  (#FilteredWorkflows)', level=2)
body(
    'The LatestVersions CTE now includes the JOIN to frs_def_workflow_type and the %form / '
    'name filters — ROW_NUMBER() only runs over matching workflows, not the entire definition '
    'table. WITH (NOLOCK) on all base table reads prevents blocking the live Ivanti system. '
    'The workflow XML is cleaned of its declaration and namespace prefix, cast to the XML data '
    'type, and stored with a UPPER()-normalised RecID. A clustered index on '
    'WorkflowDefinitionRecID is created after population.'
)

heading('Stage 2 — Single XML Shred Pass  (#AllBlocks)', level=2)
body(
    'All workflow XML is shredded into one row per block in a single pass using '
    'CROSS APPLY ... .nodes(). The block XML fragment is stored alongside title and type. '
    'DISTINCT is intentionally omitted — the XML type is not comparable in SQL Server, so '
    'DISTINCT on an XML column would error. Deduplication is handled in the PATH steps which '
    'do not carry the XML column forward. Two indexes are created: a clustered index on '
    'WorkflowDefinitionRecID and a non-clustered index on BlockType so PATH 3 can seek '
    'directly to vote0007/vote rows without scanning the full table.'
)

heading('PATH 1 — QuickAction Blocks  (#Blocks)', level=2)
body(
    'Extracts blocks with a QuickAction property and captures the QAID GUID as a '
    'uniqueidentifier. Covers advancedtask, update, create, notification, quickaction, and '
    'createnew0002. The team name is resolved in the final SELECT via CHARINDEX string scanning '
    'on frs_def_quick_actions.Definition (see Final UNION section).'
)
code_block(
    "SELECT DISTINCT ... TRY_CAST(\n"
    "    q.qaprop.value('(groups/group/param[name=\"QAID\"]/value)[1]', 'nvarchar(100)')\n"
    "AS uniqueidentifier) AS QAID\n"
    "INTO #Blocks\n"
    "FROM #AllBlocks ab\n"
    "CROSS APPLY ab.BlockXml.nodes('block/blockProperties/property[name=\"QuickAction\"]') q(qaprop);\n"
    "CREATE CLUSTERED INDEX    IX_Blocks_QAID  ON #Blocks (QAID);\n"
    "CREATE NONCLUSTERED INDEX IX_Blocks_RecID ON #Blocks (WorkflowDefinitionRecID);"
)

heading('PATH 2 — Task Blocks  (#TaskBlocks)', level=2)
body(
    'Extracts task blocks via the teamblock XML property. CROSS APPLY (VALUES) computes '
    'TeamName once to avoid repeated XPath evaluation. Blank values and dynamic expressions '
    '(starting with "$(") are filtered out.'
)

heading('Pre-stage — Approval Group Lookup  (#ApprovalGroupLookup)', level=2)
body(
    'Before PATH 3, active Service Request Approval contact groups are pre-materialised into '
    '#ApprovalGroupLookup with a clustered index on RecId. This avoids PATH 3 scanning the '
    'full ContactGroup table on every query.'
)
code_block(
    "SELECT RecId, Name INTO #ApprovalGroupLookup\n"
    "FROM ContactGroup WITH (NOLOCK)\n"
    "WHERE Status = 'Active' AND GroupType = 'Service Request Approval';\n"
    "CREATE CLUSTERED INDEX IX_AGL_RecId ON #ApprovalGroupLookup (RecId);"
)

heading('PATH 3 — Approval Blocks  (#ApprovalBlocks)', level=2)
body(
    'Extracts vote0007 and vote approval blocks. The approver contact group GUID is read from '
    'the approvers/groups/group XML path, then joined to #ApprovalGroupLookup to resolve the '
    'GUID to a group name. UPPER() is applied to the XML-extracted value (casing not guaranteed) '
    'but removed from the JOIN column so the clustered index on RecId is sargable on CI collations. '
    'Only static group assignments (where a GUID is present) are included; dynamic approvers '
    'produce no GUID and are excluded by the JOIN.'
)
code_block(
    "CROSS APPLY (VALUES (\n"
    "    UPPER(LTRIM(RTRIM(g.grp.value('(param[name=\"contactgroup\"]/value)[1]', 'nvarchar(50)'))))\n"
    ")) av (ContactGroupId)\n"
    "JOIN #ApprovalGroupLookup cg ON cg.RecId = av.ContactGroupId\n"
    "WHERE ab.BlockType IN ('vote0007', 'vote') AND av.ContactGroupId <> '';"
)

heading('Stage 3 — Materialise Offering Status  (#WorkflowOffering)', level=2)
body(
    'The three-table offering status join (ServiceReqFulfillmentPlan → FusionLink → '
    'ServiceReqTemplate) is materialised once into a temp table with a clustered index. '
    'All three UNION branches read from this table. WITH (NOLOCK) applied to all three '
    'base tables.'
)

heading('Final UNION Result', level=2)
body('Three branches joined to #WorkflowOffering for offering status:')
bullet(
    '#Blocks branch (PATH 1): joins to frs_def_quick_actions WITH (NOLOCK) and uses '
    'CHARINDEX string scanning to extract the OwnerTeam value from the Definition JSON column. '
    'OPENJSON cannot be used because Ivanti stores JavaScript Date literals (new Date(...)) '
    'in this column. Two-step null guard: checks etPos.pos > 0 before extracting so null '
    'ExpressionText values return NULL rather than arbitrary text.'
)
bullet('#TaskBlocks branch (PATH 2): team name already in #TaskBlocks, no further lookup.')
bullet('#ApprovalBlocks branch (PATH 3): group name already resolved via #ApprovalGroupLookup JOIN.')

# ── 5. Changes from v7 to v8 ─────────────────────────────────────────────────
heading('5. Changes from v7 to v8')
add_table(
    ['#', 'Change', 'Reason', 'Benefit'],
    [
        ['1',
         'PATH 3 added: vote0007 and vote approval blocks extracted via ContactGroup JOIN',
         'These block types carry approver group assignments not captured by PATH 1 or PATH 2',
         'Approval blocks now appear in results alongside task and QuickAction blocks'],
        ['2',
         'PATH 1 comment updated to explicitly list all covered block types: advancedtask, update, create, notification, quickaction, createnew0002',
         'All six types carry a QuickAction/QAID property and were already captured; the coverage was implicit and undocumented',
         'Clarity — no SQL change; makes the intent explicit'],
        ['3',
         '#ApprovalBlocks added to the DROP TABLE cleanup section',
         'Temp table cleanup consistency',
         'No orphaned temp tables if the script is run multiple times in the same session'],
        ['4',
         'Filter pushed into LatestVersions CTE — JOIN to frs_def_workflow_type and %form name filter now inside the CTE',
         'ROW_NUMBER() was running over the entire definition table before filtering',
         'ROW_NUMBER() now only runs over matching *form workflows'],
        ['5',
         'Non-clustered index on #AllBlocks(BlockType)',
         'PATH 3 was doing a full scan of #AllBlocks to find vote0007/vote rows',
         'PATH 3 can now seek directly to the relevant rows'],
        ['6',
         '#ApprovalGroupLookup pre-materialises active Service Request Approval groups before PATH 3',
         'PATH 3 was joining the full ContactGroup table on every execution',
         'PATH 3 joins a small indexed temp table instead'],
        ['7',
         'UPPER() removed from the JOIN column in PATH 3 (cg.RecId)',
         'UPPER() on an indexed column prevents the index from being used (not sargable)',
         'On CI collations the clustered index on RecId is now usable for the JOIN'],
        ['8',
         'Non-clustered index on #Blocks(WorkflowDefinitionRecID)',
         'The final LEFT JOIN to #WorkflowOffering in PATH 1 had no index support on the #Blocks side',
         'Faster LEFT JOIN in the final SELECT'],
    ]
)

# ── 6. Database Tables Referenced ────────────────────────────────────────────
heading('6. Database Tables Referenced')
add_table(
    ['Table', 'Purpose'],
    [
        ['frs_def_workflow_definition',  'Stores versioned workflow XML definitions'],
        ['frs_def_workflow_type',        'Stores workflow names and types'],
        ['frs_def_quick_actions',        'Stores Quick Action JSON definitions (contains OwnerTeam assignments)'],
        ['ServiceReqFulfillmentPlan',    'Links workflows to service request fulfillment plans'],
        ['FusionLink',                   'Relationship table linking fulfillment plans to service request templates'],
        ['ServiceReqTemplate',           'Service request offering metadata including Status'],
        ['StandardUserTeam',             'Active service desk teams (used by the web app dropdown)'],
        ['ContactGroup',                 'Contact groups; approval groups filtered by GroupType = "Service Request Approval"'],
    ]
)

# ── 7. Output Columns ─────────────────────────────────────────────────────────
heading('7. Output Columns')
add_table(
    ['Column', 'Description'],
    [
        ['WorkflowName',          'Name of the workflow'],
        ['DefVersion',            'Version number of the workflow definition'],
        ['RequestOfferingStatus', 'Status of the associated request offering (Published or Design). Workflows not linked to a request offering are excluded from results.'],
        ['BlockTitle',            'Display title of the block within the workflow'],
        ['BlockType',             'Type of block (see Supported Block Types)'],
        ['TeamName',              'Team assigned to the block, or approval group name for vote0007/vote blocks'],
    ]
)

# ── 8. Sample Results ─────────────────────────────────────────────────────────
heading('8. Sample Results  (Team = "Risk Management Support", Status = "Published")')
add_table(
    ['WorkflowName', 'Ver', 'Offering Status', 'Block Title', 'Block Type', 'Team / Group'],
    [
        ['Audit Data Request Request form',                         '11', 'Published (Automatic)', 'Set Team and Classification',                'update',  'Risk Management Support'],
        ['Audit Data Request Request form',                         '11', 'Published (Automatic)', 'Set Team and Classification 2',              'update',  'Risk Management Support'],
        ['Request Software Not in the Company Portal Request form', '12', 'Published (Automatic)', '1 - Fulfill Risk Assessment Request Lite',  'task',    'Risk Management Support'],
        ['Request Software Not in the Company Portal Request form', '12', 'Published (Automatic)', '1a - Fulfill Risk Assessment Request',      'task',    'Risk Management Support'],
    ]
)

# ── 9. Common Use Cases ───────────────────────────────────────────────────────
heading('9. Common Use Cases')
bullet('Audit which workflows route work to a specific team before renaming or retiring that team.')
bullet('Identify all task and update blocks assigned to a team across all published workflows.')
bullet('Find all approval blocks assigned to a specific contact group across all workflows.')
bullet('Verify workflow version alignment — confirm the latest version is what is actively used.')
bullet('Pre-migration checklist — find all workflows referencing a team being restructured.')
bullet('Set @TeamName and @Status to blank for a full cross-team and cross-group routing map.')

# ── 10. Tips & Customization ───────────────────────────────────────────────────
heading('10. Tips & Customization')
code_block(
    "DECLARE @BlockType NVARCHAR(50) = 'vote0007'; -- only approval blocks\n"
    "DECLARE @TeamName  NVARCHAR(255) = 'Legal';   -- partial match on group name"
)
bullet("Set @BlockType to 'vote0007' or 'vote' to see only approval blocks.")
bullet("Set @BlockType to 'task' to see only task assignments.")
bullet('Leave @TeamName blank to generate a full routing map across all teams and approval groups.')
bullet("Change @Status to 'Design' to audit workflows still in development.")
bullet("Workflow filter targets names ending in 'form' — adjust LIKE '%form' to broaden scope.")

# ── 11. Dependencies & Notes ──────────────────────────────────────────────────
heading('11. Dependencies & Notes')
bullet('Requires SELECT access on: frs_def_workflow_definition, frs_def_workflow_type, frs_def_quick_actions, ServiceReqFulfillmentPlan, FusionLink, ServiceReqTemplate, ContactGroup.')
bullet('Uses temporary tables and CREATE INDEX — requires tempdb write access.')
bullet('Temp tables are cleaned up at the end via DROP TABLE statements.')
bullet('XML namespace and declaration are stripped before parsing — required because SQL Server rejects the Ivanti-specific namespace URI.')
bullet(
    'Quick Action team extraction uses CHARINDEX string scanning on Definition. '
    'OPENJSON cannot be used because Ivanti stores JavaScript Date literals (new Date(...)) '
    'in this column, which are not valid JSON.'
)
bullet('Approval blocks with dynamic approvers (_unchecked, from field, from profile) are excluded — they produce no contactgroup GUID and are filtered by the ContactGroup JOIN.')
bullet('Targets SQL Server 2016+ due to XQuery, TRY_CAST, and window function usage. Tested on SQL Server 2019.')

# ── 12. Web Application ───────────────────────────────────────────────────────
heading('12. Web Application')
body(
    'The Next.js web app (workflow-query-app) provides a browser-based interface over the '
    'v8 SQL query. Source code: https://github.com/gsecrest/workflow-query-app.'
)

heading('12.1 Filters', level=2)
bullet('Workflow Name: Partial-match text input.')
bullet('Block Type: Dropdown with all 9 supported block types.')
bullet('Team / Group: Searchable combobox — type to filter the list instantly, or click the chevron to browse all options. Shows "Service Desk Teams" (from StandardUserTeam) and "Approval Groups" (from ContactGroup where GroupType = \'Service Request Approval\') in separate sections. A ✕ button clears the selection.')
bullet('Status: Dropdown — All statuses (Published + Design), Published only, Design only. Workflows not linked to a request offering are always excluded.')

heading('12.2 Results Table', level=2)
body(
    'Results render in a table with the same six columns. '
    'vote0007 and vote rows show a small "group" tag inline next to the name in the '
    'Team / Group column to distinguish approval groups from service desk teams at a glance.'
)

heading('12.3 Exporting Results', level=2)
bullet('Export CSV — Downloads results as workflow-results.csv with UTF-8 BOM for Excel compatibility.')
bullet('Copy to Clipboard — Copies results as tab-separated values for direct paste into Excel.')

heading('12.4 API Routes', level=2)
add_table(
    ['Route', 'Method', 'Description'],
    [
        ['/api/query',                        'POST', 'Accepts { workflowName, blockType, teamName, status } and returns JSON rows.'],
        ['/api/teams',                        'GET',  'Returns { teams, approvalGroups } — service desk teams and approval contact groups.'],
        ['/api/export/workflow-results.csv',  'GET',  'Re-runs the query with filter params and returns a CSV download.'],
    ]
)

# ── 13. Repository Documentation ─────────────────────────────────────────────
heading('13. Repository Documentation')
body(
    'SQL scripts: https://github.com/gsecrest/ivanti-workflow-tools  '
    'Web app: https://github.com/gsecrest/workflow-query-app'
)
heading('13.1 Version History', level=2)
add_table(
    ['File', 'Notes'],
    [
        ['FindTeamByBlockTypeAndWorkflow.sql',    'v1 — initial version'],
        ['FindTeamByBlockTypeAndWorkflow_v2.sql', 'v2'],
        ['FindTeamByBlockTypeAndWorkflow_v3.sql', 'v3'],
        ['FindTeamByBlockTypeAndWorkflow_v4.sql', 'v4 — see v4 documentation'],
        ['FindTeamByBlockTypeAndWorkflow_v5.sql', 'v5 — performance-optimised rewrite'],
        ['FindTeamByBlockTypeAndWorkflow_v6.sql', 'v6 — single XML shred pass; WorkflowOffering as temp table'],
        ['FindTeamByBlockTypeAndWorkflow_v7.sql', 'v7 — NOLOCK hints, DISTINCT fix on #AllBlocks, improved CHARINDEX null guard'],
        ['FindTeamByBlockTypeAndWorkflow_v8.sql', 'v8 — PATH 3 for approval blocks; expanded QuickAction block type coverage (current)'],
    ]
)

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/Glenn.Secrest/Downloads/Claude/SQL/FindTeamByBlockTypeAndWorkflow_v8_Documentation.docx'
doc.save(out)
print(f'Saved: {out}')
