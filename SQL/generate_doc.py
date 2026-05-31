from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def body(text):
    return doc.add_paragraph(text)

def bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0x7C)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'EEF2F8')
    p._p.get_or_add_pPr().append(shading)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1F3864')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        hdr[i]._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_heading('FindTeamByBlockTypeAndWorkflow_v4.sql', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph('Technical Documentation — Ivanti ITSM Workflow Analysis Query')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
sub.runs[0].font.italic = True

doc.add_paragraph()

# ── 1. Overview ──────────────────────────────────────────────────────────────
heading('1. Overview')
body(
    'This SQL script is a diagnostic and audit tool for the Ivanti ITSM (Neurons) platform. '
    'It inspects published workflow definitions and identifies which teams are assigned within '
    'specific workflow blocks — either via Quick Actions (update/advancedtask blocks) or direct '
    'task block assignments. The results help administrators understand team routing logic '
    'embedded in service request workflows without navigating the Ivanti UI.'
)

# ── 2. Input Parameters ───────────────────────────────────────────────────────
heading('2. Input Parameters')
body('Four DECLARE variables at the top of the script control filtering. Leave any blank to return all matching records.')
doc.add_paragraph()
add_table(
    ['Parameter', 'Type', 'Default Value', 'Description'],
    [
        ['@WorkflowName', 'NVARCHAR(255)', "(blank — all)", 'Filter by workflow name (partial match). Leave blank to search all workflows.'],
        ['@BlockType',    'NVARCHAR(50)',  "(blank — all)", 'Filter by block type: task, advancedtask, or update. Leave blank for all types.'],
        ['@TeamName',     'NVARCHAR(255)', 'Risk Management Support', 'Filter by team name (partial match). Leave blank to return all teams.'],
        ['@Status',       'NVARCHAR(50)',  'Published', 'Filter by request offering status (e.g. Published, Design). Leave blank for all.'],
    ]
)

# ── 3. How It Works ───────────────────────────────────────────────────────────
heading('3. How It Works')
body('The script executes in four logical phases:')

heading('Phase 1 — Load Filtered Workflows  (temp table: #FilteredWorkflows)', level=2)
body(
    'Reads from frs_def_workflow_definition and frs_def_workflow_type to find the most recent '
    'version of every workflow whose name ends in "form" (excluding backups). '
    'The workflow XML stored in the Details column is cleaned of its XML declaration and '
    'namespace prefix before being cast to the SQL XML data type for parsing in later steps.'
)
code_block('SELECT wt.Name, wf.RecID, wf.DefVersion, CAST(... AS XML) AS XmlData\nINTO #FilteredWorkflows\nFROM frs_def_workflow_definition wf\nJOIN frs_def_workflow_type wt ...\nWHERE wt.Name LIKE \'%form\' AND CAST(wf.DefVersion AS INT) = (MAX version subquery)')

heading('Phase 2 — Extract Quick Action Blocks  (temp table: #Blocks)', level=2)
body(
    'Uses XQuery (CROSS APPLY / .nodes() / .value()) to shred the workflow XML and extract '
    'blocks that contain a QuickAction property. These are typically update or advancedtask '
    'blocks that trigger a Quick Action to set field values. The QAID (Quick Action ID) is '
    'captured for later lookup.'
)
code_block('CROSS APPLY fw.XmlData.nodes(\'/scenario/blocks/block\') b(block)\nCROSS APPLY b.block.nodes(\'blockProperties/property[name="QuickAction"]\') q(qaprop)')

heading('Phase 3 — Extract Task Blocks  (temp table: #TaskBlocks)', level=2)
body(
    'Parses the same XML for blocks that contain a teamblock property — these are standard '
    'task blocks where a team is assigned directly. Filters out blank values and '
    'dynamic expressions (values starting with "$(") to return only static team names.'
)
code_block('CROSS APPLY b.block.nodes(\'blockProperties/property[name="teamblock"]/groups/group/param[name="team"]\') p(prop)')

heading('Phase 4 — Final UNION Result', level=2)
body(
    'Combines both paths into a single result set. For Quick Action blocks, it looks up '
    'the Quick Action definition JSON in frs_def_quick_actions and extracts the OwnerTeam '
    'value using CHARINDEX string parsing. For task blocks, the team name is already '
    'available from Phase 3. Both paths join to ServiceReqFulfillmentPlan, FusionLink, '
    'and ServiceReqTemplate to enrich results with the associated request offering status.'
)

# ── 4. Tables Referenced ──────────────────────────────────────────────────────
heading('4. Database Tables Referenced')
add_table(
    ['Table', 'Purpose'],
    [
        ['frs_def_workflow_definition',  'Stores versioned workflow XML definitions'],
        ['frs_def_workflow_type',        'Stores workflow names/types'],
        ['frs_def_quick_actions',        'Stores Quick Action JSON definitions (contains OwnerTeam field assignments)'],
        ['ServiceReqFulfillmentPlan',    'Links workflows to service request fulfillment plans'],
        ['FusionLink',                   'Relationship table linking fulfillment plans to service request templates'],
        ['ServiceReqTemplate',           'Service request offering metadata including Status (Published, Design, etc.)'],
    ]
)

# ── 5. Output Columns ─────────────────────────────────────────────────────────
heading('5. Output Columns')
add_table(
    ['Column', 'Description'],
    [
        ['WorkflowName',          'Name of the workflow (e.g. "Audit Data Request Request form")'],
        ['DefVersion',            'Version number of the workflow definition'],
        ['RequestOfferingStatus', 'Status of the associated request offering (Published, Design, or "No Offering" if unlinked)'],
        ['BlockTitle',            'Display title of the block within the workflow'],
        ['BlockType',             'Type of block: task, advancedtask, or update'],
        ['TeamName',              'Team assigned within that block'],
    ]
)

# ── 6. Sample Results ─────────────────────────────────────────────────────────
heading('6. Sample Results  (Team = "Risk Management Support", Status = "Published")')
add_table(
    ['WorkflowName', 'Ver', 'Offering Status', 'Block Title', 'Block Type', 'Team'],
    [
        ['Audit Data Request Request form',                       '11', 'Published (Automatic)', 'Set Team and Classification',               'update', 'Risk Management Support'],
        ['Audit Data Request Request form',                       '11', 'Published (Automatic)', 'Set Team and Classification 2',             'update', 'Risk Management Support'],
        ['Request Software Not in the Company Portal Request form','12', 'Published (Automatic)', '1 - Fulfill Risk Assessment Request Lite', 'task',   'Risk Management Support'],
        ['Request Software Not in the Company Portal Request form','12', 'Published (Automatic)', '1a - Fulfill Risk Assessment Request',     'task',   'Risk Management Support'],
    ]
)

# ── 7. Common Use Cases ───────────────────────────────────────────────────────
heading('7. Common Use Cases')
bullet('Audit which workflows route work to a specific team before renaming or retiring that team.')
bullet('Identify all task and update blocks assigned to a team across all published workflows.')
bullet('Verify workflow version alignment — ensure the latest version is what is actively used.')
bullet('Pre-migration checklist — find all workflows that reference a team being restructured.')
bullet('Change all @TeamName and @Status to blank to get a full cross-team routing map.')

# ── 8. Tips & Customization ───────────────────────────────────────────────────
heading('8. Tips & Customization')
body('Adjust the parameters at the top of the script to change scope:')
code_block(
    "DECLARE @WorkflowName NVARCHAR(255) = '';          -- blank = all workflows\n"
    "DECLARE @BlockType    NVARCHAR(50)  = 'task';      -- only task blocks\n"
    "DECLARE @TeamName     NVARCHAR(255) = '';          -- blank = all teams\n"
    "DECLARE @Status       NVARCHAR(50)  = 'Published'; -- only published offerings"
)
bullet('Change @BlockType to "update" to see only Quick Action assignment blocks.')
bullet('Leave @TeamName blank to generate a full routing map across all teams.')
bullet('Change @Status to "Design" to audit workflows still in development.')
bullet('The WHERE clause filters workflows ending in "form" — adjust LIKE \'%form\' to broaden scope.')

# ── 9. Dependencies & Notes ───────────────────────────────────────────────────
heading('9. Dependencies & Notes')
bullet('Requires SELECT access on: frs_def_workflow_definition, frs_def_workflow_type, frs_def_quick_actions, ServiceReqFulfillmentPlan, FusionLink, ServiceReqTemplate.')
bullet('Uses temporary tables (#FilteredWorkflows, #Blocks, #TaskBlocks) — requires tempdb write access.')
bullet('Temp tables are cleaned up at the end of the script via DROP TABLE statements.')
bullet('The XML namespace and declaration are stripped manually before parsing — required because SQL Server\'s XML parser rejects the Ivanti-specific namespace URI.')
bullet('Quick Action team extraction uses string parsing (CHARINDEX) on JSON — fragile if the JSON schema changes in future Ivanti versions.')
bullet('The script targets SQL Server 2016+ due to XQuery and TRY_CAST usage.')

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/Glenn.Secrest/Downloads/Claude/SQL/FindTeamByBlockTypeAndWorkflow_v4_Documentation.docx'
doc.save(out)
print(f'Saved: {out}')
