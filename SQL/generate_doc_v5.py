from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def heading(text, level=1):
    return doc.add_heading(text, level=level)

def body(text):
    return doc.add_paragraph(text)

def bullet(text):
    return doc.add_paragraph(text, style='List Bullet')

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1E, 0x40, 0x7C)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'), 'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'), 'EEF2F8')
    p._p.get_or_add_pPr().append(shading)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        run = hdr[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = OxmlElement('w:shd')
        shading.set(qn('w:val'), 'clear')
        shading.set(qn('w:color'), 'auto')
        shading.set(qn('w:fill'), '1F3864')
        hdr[i]._tc.get_or_add_tcPr().append(shading)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = val
    doc.add_paragraph()

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_heading('FindTeamByBlockTypeAndWorkflow_v5.sql', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph('Technical Documentation — Ivanti ITSM Workflow Analysis Query (Optimised)')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.color.rgb = RGBColor(0x44, 0x44, 0x44)
sub.runs[0].font.italic = True
doc.add_paragraph()

# ── 1. Overview ───────────────────────────────────────────────────────────────
heading('1. Overview')
body(
    'This SQL script is a diagnostic and audit tool for the Ivanti ITSM (Neurons) platform. '
    'It inspects published workflow definitions and identifies which teams are assigned within '
    'specific workflow blocks — either via Quick Actions (update/advancedtask blocks) or direct '
    'task block assignments. The results help administrators understand team routing logic '
    'embedded in service request workflows without navigating the Ivanti UI.'
)
body(
    'v5 is a performance-optimised rewrite of v4. All output columns and result sets are '
    'identical; only the internal execution strategy has changed.'
)

# ── 2. Input Parameters ───────────────────────────────────────────────────────
heading('2. Input Parameters')
body('Four DECLARE variables at the top of the script control filtering. Leave any blank to return all matching records.')
doc.add_paragraph()
add_table(
    ['Parameter', 'Type', 'Default Value', 'Description'],
    [
        ['@WorkflowName', 'NVARCHAR(255)', '(blank — all)',           'Filter by workflow name (partial match).'],
        ['@BlockType',    'NVARCHAR(50)',  '(blank — all)',           'Filter by block type: task, advancedtask, or update.'],
        ['@TeamName',     'NVARCHAR(255)', 'Risk Management Support', 'Filter by team name (partial match).'],
        ['@Status',       'NVARCHAR(50)',  'Published',               'Filter by request offering status (e.g. Published, Design).'],
    ]
)

# ── 3. How It Works ───────────────────────────────────────────────────────────
heading('3. How It Works')
body('The script executes in four logical phases:')

heading('Phase 1 — Load Filtered Workflows  (#FilteredWorkflows)', level=2)
body(
    'A CTE (LatestVersions) uses ROW_NUMBER() partitioned by WorkflowTypeLink_RecID to identify '
    'the most recent version of each workflow in a single pass. This replaces the v4 correlated '
    'subquery that re-executed MAX() once per outer row. The workflow XML is cleaned of its '
    'declaration and namespace prefix, cast to the XML data type, and stored alongside a '
    'UPPER()-normalised RecID to avoid repeated UPPER() calls in later joins. A clustered index '
    'on WorkflowDefinitionRecID is created immediately after population.'
)
code_block(
    ";WITH LatestVersions AS (\n"
    "    SELECT RecID, WorkflowTypeLink_RecID, DefVersion, Details,\n"
    "           ROW_NUMBER() OVER (\n"
    "               PARTITION BY WorkflowTypeLink_RecID\n"
    "               ORDER BY CAST(DefVersion AS INT) DESC\n"
    "           ) AS rn\n"
    "    FROM frs_def_workflow_definition\n"
    ")\n"
    "SELECT ... INTO #FilteredWorkflows FROM LatestVersions WHERE rn = 1 ...;\n"
    "CREATE CLUSTERED INDEX IX_FW_RecID ON #FilteredWorkflows (WorkflowDefinitionRecID);"
)

heading('Phase 2 — Extract Quick Action Blocks  (#Blocks)', level=2)
body(
    'Uses XQuery (CROSS APPLY / .nodes() / .value()) to shred the workflow XML and extract '
    'blocks containing a QuickAction property. The QAID (Quick Action ID) is captured for '
    'later lookup in frs_def_quick_actions. A clustered index on QAID is created after '
    'population to speed up the JOIN in the final query.'
)
code_block(
    "CROSS APPLY fw.XmlData.nodes('/scenario/blocks/block') b(block)\n"
    "CROSS APPLY b.block.nodes('blockProperties/property[name=\"QuickAction\"]') q(qaprop)\n"
    "...\n"
    "CREATE CLUSTERED INDEX IX_Blocks_QAID ON #Blocks (QAID);"
)

heading('Phase 3 — Extract Task Blocks  (#TaskBlocks)', level=2)
body(
    'Parses the same XML for blocks that contain a teamblock property. Filters out blank '
    'values and dynamic expressions (starting with "$("). A clustered index on '
    'WorkflowDefinitionRecID is created after population.'
)
code_block(
    "CROSS APPLY b.block.nodes('blockProperties/property[name=\"teamblock\"]"
    "/groups/group/param[name=\"team\"]') p(prop)\n"
    "...\n"
    "CREATE CLUSTERED INDEX IX_TaskBlocks_RecID ON #TaskBlocks (WorkflowDefinitionRecID);"
)

heading('Phase 4 — Final UNION Result', level=2)
body(
    'A WorkflowOffering CTE pre-computes the three-table offering status join '
    '(ServiceReqFulfillmentPlan → FusionLink → ServiceReqTemplate) once. In v4 this join '
    'was duplicated in both UNION branches. UPPER() is applied to WorkflowId inside the CTE '
    'so the join condition in each branch is a direct string comparison.'
)
body(
    'For Quick Action blocks, the team name is extracted from the Quick Action JSON using '
    'three chained CROSS APPLY (VALUES(...)) steps that build the CHARINDEX/SUBSTRING '
    'calculation incrementally. In v4 the full expression was repeated four times across '
    'SELECT and WHERE. The alias tn.TeamName is referenced once in each clause.'
)
code_block(
    ";WITH WorkflowOffering AS (\n"
    "    SELECT DISTINCT UPPER(fp.WorkflowId) AS WorkflowId, srt.Status\n"
    "    FROM ServiceReqFulfillmentPlan fp\n"
    "    JOIN FusionLink fl   ON fl.TargetID = fp.RecId AND fl.RelationshipName = '...'\n"
    "    JOIN ServiceReqTemplate srt ON srt.RecId = fl.SourceID\n"
    ")\n"
    "...\n"
    "CROSS APPLY (VALUES (CHARINDEX('\"FieldName\":\"OwnerTeam\"', qa.Definition))) ownerPos(pos)\n"
    "CROSS APPLY (VALUES (CHARINDEX('\"ExpressionText\":\"', qa.Definition, ownerPos.pos) + 18)) valStart(idx)\n"
    "CROSS APPLY (VALUES (\n"
    "    CASE WHEN ownerPos.pos > 0\n"
    "         THEN LEFT(SUBSTRING(qa.Definition, valStart.idx, 500),\n"
    "                   CHARINDEX('\"', SUBSTRING(qa.Definition, valStart.idx, 500)) - 1)\n"
    "    END\n"
    ")) tn(TeamName)"
)

# ── 4. v4 → v5 Changes ───────────────────────────────────────────────────────
heading('4. Changes from v4 to v5')
add_table(
    ['#', 'Problem in v4', 'Fix in v5', 'Benefit'],
    [
        ['1', 'Correlated MAX subquery ran once per row in frs_def_workflow_definition',
              'ROW_NUMBER() OVER (PARTITION BY WorkflowTypeLink_RecID)',
              'Single table scan instead of N subquery executions'],
        ['2', 'No indexes on temp tables',
              'CREATE CLUSTERED INDEX after each SELECT INTO',
              'Faster JOIN lookups in the final query'],
        ['3', '3-table offering JOIN duplicated in both UNION branches',
              'Single WorkflowOffering CTE referenced by both branches',
              'Offering join executes once instead of twice'],
        ['4', 'UPPER() applied to both sides of every JOIN condition',
              'UPPER() applied once at storage time (#FilteredWorkflows) and once in the CTE',
              'Eliminates per-row function call on join columns'],
        ['5', 'TeamName CHARINDEX/SUBSTRING expression repeated 4× in SELECT and WHERE',
              'Three CROSS APPLY (VALUES(...)) steps compute the value once as tn.TeamName',
              'Single evaluation; easier to maintain if JSON structure changes'],
    ]
)

# ── 5. Database Tables Referenced ────────────────────────────────────────────
heading('5. Database Tables Referenced')
add_table(
    ['Table', 'Purpose'],
    [
        ['frs_def_workflow_definition',  'Stores versioned workflow XML definitions'],
        ['frs_def_workflow_type',        'Stores workflow names and types'],
        ['frs_def_quick_actions',        'Stores Quick Action JSON definitions (contains OwnerTeam assignments)'],
        ['ServiceReqFulfillmentPlan',    'Links workflows to service request fulfillment plans'],
        ['FusionLink',                   'Relationship table linking fulfillment plans to service request templates'],
        ['ServiceReqTemplate',           'Service request offering metadata including Status'],
    ]
)

# ── 6. Output Columns ─────────────────────────────────────────────────────────
heading('6. Output Columns')
add_table(
    ['Column', 'Description'],
    [
        ['WorkflowName',          'Name of the workflow (e.g. "Audit Data Request Request form")'],
        ['DefVersion',            'Version number of the workflow definition'],
        ['RequestOfferingStatus', 'Status of the associated request offering, or "No Offering" if unlinked'],
        ['BlockTitle',            'Display title of the block within the workflow'],
        ['BlockType',             'Type of block: task, advancedtask, or update'],
        ['TeamName',              'Team assigned within that block'],
    ]
)

# ── 7. Sample Results ─────────────────────────────────────────────────────────
heading('7. Sample Results  (Team = "Risk Management Support", Status = "Published")')
add_table(
    ['WorkflowName', 'Ver', 'Offering Status', 'Block Title', 'Block Type', 'Team'],
    [
        ['Audit Data Request Request form',                        '11', 'Published (Automatic)', 'Set Team and Classification',               'update', 'Risk Management Support'],
        ['Audit Data Request Request form',                        '11', 'Published (Automatic)', 'Set Team and Classification 2',             'update', 'Risk Management Support'],
        ['Request Software Not in the Company Portal Request form','12', 'Published (Automatic)', '1 - Fulfill Risk Assessment Request Lite', 'task',   'Risk Management Support'],
        ['Request Software Not in the Company Portal Request form','12', 'Published (Automatic)', '1a - Fulfill Risk Assessment Request',     'task',   'Risk Management Support'],
    ]
)

# ── 8. Common Use Cases ───────────────────────────────────────────────────────
heading('8. Common Use Cases')
bullet('Audit which workflows route work to a specific team before renaming or retiring that team.')
bullet('Identify all task and update blocks assigned to a team across all published workflows.')
bullet('Verify workflow version alignment — confirm the latest version is what is actively used.')
bullet('Pre-migration checklist — find all workflows referencing a team being restructured.')
bullet('Set @TeamName and @Status to blank for a full cross-team routing map.')

# ── 9. Tips & Customization ───────────────────────────────────────────────────
heading('9. Tips & Customization')
code_block(
    "DECLARE @WorkflowName NVARCHAR(255) = '';          -- blank = all workflows\n"
    "DECLARE @BlockType    NVARCHAR(50)  = 'task';      -- only task blocks\n"
    "DECLARE @TeamName     NVARCHAR(255) = '';          -- blank = all teams\n"
    "DECLARE @Status       NVARCHAR(50)  = 'Published'; -- only published offerings"
)
bullet("Change @BlockType to 'update' to see only Quick Action assignment blocks.")
bullet('Leave @TeamName blank to generate a full routing map across all teams.')
bullet("Change @Status to 'Design' to audit workflows still in development.")
bullet("Workflow filter targets names ending in 'form' — adjust LIKE '%form' to broaden scope.")

# ── 10. Dependencies & Notes ──────────────────────────────────────────────────
heading('10. Dependencies & Notes')
bullet('Requires SELECT access on: frs_def_workflow_definition, frs_def_workflow_type, frs_def_quick_actions, ServiceReqFulfillmentPlan, FusionLink, ServiceReqTemplate.')
bullet('Uses temporary tables and CREATE INDEX — requires tempdb write access.')
bullet('Temp tables are cleaned up at the end via DROP TABLE statements.')
bullet('XML namespace and declaration are stripped before parsing — required because SQL Server rejects the Ivanti-specific namespace URI.')
bullet('Quick Action team extraction uses CHARINDEX string parsing on JSON — verify behaviour if the Quick Action JSON schema changes in a future Ivanti release.')
bullet('Targets SQL Server 2016+ due to XQuery, TRY_CAST, and window function usage. Tested on SQL Server 2019.')

# ── Save ──────────────────────────────────────────────────────────────────────
out = '/Users/Glenn.Secrest/Downloads/Claude/SQL/FindTeamByBlockTypeAndWorkflow_v5_Documentation.docx'
doc.save(out)
print(f'Saved: {out}')
